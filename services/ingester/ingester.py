"""
ingester — walk CORPUS_PATH, route by file type, chunk per docs/ingestion-pipeline.md,
embed (dense + sparse) via the step-03 embedder, upsert into Qdrant with named
vectors for hybrid retrieval.

Reference implementation against docs/ingestion-pipeline.md. CPU embedder (BGE-M3),
NOT the NPU. See ADR-0002.

Routing:
    .xlsx/.xls               -> table path (services/ingester/xlsx_tables.py)
    .csv                     -> table path (single-header CSV)
    .docx/.txt/.md/.markdown -> prose path (semantic chunks w/ overlap)
    .pdf                     -> prose path (text-layer extraction; scanned PDFs
                                with no text layer are skipped with a warning)
    else                     -> skip with warning
Vectors:    named 'dense' (1024, cosine) + named sparse 'sparse'
Point IDs are opaque uuid4; idempotency + pruning live in the checksum manifest
(corpus_manifest.py, ADR-0006), NOT in the point ID. Modes: default = in-place
incremental sync; --recreate = full rebuild + atomic alias switch; --watch = a
debounced autonomous watcher. Retrieval targets the ALIAS, not a physical collection.

ALL messy-XLSX structural handling (multi-table sheets, stacked headers, multi-row
title+band headers, carry-forward and vertical-merge grouping, headerless data,
trailing totals/notes trimming) now lives in xlsx_tables.py. This file keeps prose
routing, embedding, and Qdrant upsert.
"""
import os
import re
import sys
import csv
import time
import uuid
import hashlib
import logging
import pathlib
import argparse
from typing import Iterator

import requests
from qdrant_client import QdrantClient
from qdrant_client import models as qm

from xlsx_tables import iter_xlsx, point_id, clean_cell, _serialize_kv
from corpus_manifest import Manifest, pass_lock, Locked

# --- config -----------------------------------------------------------------
QDRANT_URL = os.environ.get("QDRANT_URL", "http://localhost:6333")
EMBED      = os.environ.get("EMBED_ENDPOINT", "http://localhost:8090/embed")
ROOT       = pathlib.Path(os.environ.get("CORPUS_PATH", "/data/corpus"))
ALIAS      = os.environ.get("QDRANT_COLLECTION", "corpus")   # retrieval targets this ALIAS (ADR-0006)
DENSE_DIM  = 1024
BATCH      = 64

# manifest + sync (ADR-0006)
MANIFEST_DB   = os.environ.get("MANIFEST_DB", str(pathlib.Path(__file__).with_name("manifest.db")))
GRACE         = float(os.environ.get("INGEST_DELETE_GRACE", "120"))    # watch: delay before pruning a vanished file
SCAN_INTERVAL = float(os.environ.get("INGEST_SCAN_INTERVAL", "3600"))  # watch: full-scan backstop
DEBOUNCE      = float(os.environ.get("INGEST_DEBOUNCE", "3"))          # watch: quiescence before a pass
WATCH_POLL    = os.environ.get("INGEST_WATCH_POLL", "").lower() in ("1", "true", "yes")

# Chunk sizes are measured in whitespace tokens (words) — a fast, deterministic
# proxy for model tokens (~1.3 model-tokens/word EN, more for CS). Values are
# calibrated so a chunk stays well under ~600 model tokens.
CHUNK_TARGET_WORDS = 300   # packing target when a section must be split (~380-400 model tokens)
CHUNK_MAX_WORDS    = 480   # hard ceiling (~600 model tokens): a section at/under this is kept
                           # WHOLE — never split merely to hit the target (no partial answers)
OVERLAP_FRAC       = 0.15  # sentence overlap carried between split chunks, as a fraction of target

logging.basicConfig(level=logging.INFO, format="%(levelname)s %(message)s")
log = logging.getLogger("ingester")

try:
    from docx import Document
except ImportError:
    Document = None
try:
    import pdfplumber
except ImportError:
    pdfplumber = None
try:
    from langdetect import detect as _detect
except ImportError:
    _detect = None


# --- embedding --------------------------------------------------------------
def embed(texts: list[str], kind: str = "passage") -> tuple[list[list[float]], list[dict]]:
    r = requests.post(EMBED, json={"input": texts, "kind": kind}, timeout=300)
    r.raise_for_status()
    data = r.json()
    dense = data["dense"]
    sparse = data.get("sparse", [{} for _ in texts])
    if len(dense) != len(texts) or len(sparse) != len(texts):
        raise RuntimeError(f"embedder returned {len(dense)}/{len(sparse)} for {len(texts)} inputs")
    return dense, sparse


def to_sparse_vector(sparse_map: dict) -> qm.SparseVector:
    if not sparse_map:
        return qm.SparseVector(indices=[], values=[])
    idx, val = [], []
    for k, v in sparse_map.items():
        idx.append(int(k)); val.append(float(v))
    return qm.SparseVector(indices=idx, values=val)


def detect_lang(text: str) -> str:
    if _detect is None:
        return "unknown"
    try:
        code = _detect(text)
    except Exception:
        return "unknown"
    return code if code in ("cs", "en") else "other"


# ============================================================================
# TABLE PATH (CSV stays here; XLSX delegated to xlsx_tables.iter_xlsx)
# ============================================================================
def iter_csv(path: pathlib.Path) -> Iterator[dict]:
    with path.open(newline="", encoding="utf-8", errors="ignore") as fh:
        reader = csv.reader(fh)
        rows = iter(reader)
        header = None
        for r in rows:
            if r and any(str(c).strip() for c in r):
                header = [str(c).strip() for c in r]
                break
        if not header:
            log.warning("%s: no header row, skipping", path.name)
            return
        keys = {i: h for i, h in enumerate(header) if h}
        ridx = 0
        for r in rows:
            ridx += 1
            values = {i: (r[i] if i < len(r) else None) for i in keys}
            text = _serialize_kv(keys, values, None)
            if not text:
                continue
            yield {
                "id": point_id(path.name, "csv", ridx),
                "text": text,
                "payload": {
                    "source_file": path.name, "sheet": "csv", "row_index": ridx,
                    "kind": "table_row", "lang": "n/a", "text": text,
                },
            }


# ============================================================================
# PROSE PATH
# ============================================================================
def _approx_tokens(s: str) -> int:
    """Whitespace-token (word) count — the deterministic size proxy for chunking."""
    return len(s.split())


# A section is (section_title, subsection_title, blocks); a block is a paragraph
# as (text, start_line, end_line) with 1-based source line span. Heading lines are
# consumed as boundaries and never appear in a block's text.
_MD_HEADING = re.compile(r"^(#{1,6})\s+(.*\S)\s*$")
# Sentence boundary: terminator (. ! ? … incl. CS) followed by whitespace + next
# glyph. This is the FINEST granularity we ever split at — a chunk never cuts a
# sentence mid-way, so retrieval can't hand the model a truncated instruction.
_SENT_SPLIT = re.compile(r"(?<=[.!?…])\s+(?=\S)")

# Markdown decoration to strip for the embedding-only text_clean. We keep link
# TARGET text AND url tokens (hostnames/paths matter to the sparse branch); only
# the syntactic scaffolding is removed so it stops polluting the dense vector.
_MD_CLEAN = [
    (re.compile(r"`{1,3}([^`]*)`{1,3}"),        r"\1"),   # code spans
    (re.compile(r"!?\[([^\]]*)\]\(([^)]+)\)"),  r"\1 \2"),# links/images -> text + url
    (re.compile(r"(\*\*|__)(.*?)\1"),           r"\2"),   # bold
    (re.compile(r"(?<!\w)(\*|_)(.*?)\1(?!\w)"), r"\2"),   # italic
    (re.compile(r"^\s{0,3}#{1,6}\s+", re.M),    ""),      # residual headings
    (re.compile(r"^\s{0,3}>\s?", re.M),         ""),      # blockquote markers
    (re.compile(r"^\s{0,3}([-*+]|\d+[.)])\s+", re.M), ""),# list markers
]


def _strip_md(text: str) -> str:
    """Strip Markdown syntax for embedding; keep all word content (incl. urls)."""
    for pat, repl in _MD_CLEAN:
        text = pat.sub(repl, text)
    text = text.replace("|", " ")                 # table pipes
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def _heading_titles(level: int, title: str, sec: str, sub: str) -> tuple[str, str]:
    """Fold a new heading into (section_title, subsection_title). A top-level (or
    first-seen) heading opens a fresh section and clears the subsection; a deeper
    heading becomes the subsection under the current section."""
    if level <= 1 or not sec:
        return title, ""
    return sec, title


def _overlap_tail(text: str) -> str:
    """Trailing whole sentences of `text`, up to OVERLAP_FRAC*target words — carried
    into the next chunk so a topic that straddles a split boundary appears in both."""
    budget = OVERLAP_FRAC * CHUNK_TARGET_WORDS
    keep, tok = [], 0
    for s in reversed(_SENT_SPLIT.split(text)):
        st = _approx_tokens(s)
        if keep and tok + st > budget:
            break
        keep.insert(0, s); tok += st
    return " ".join(keep).strip()


def _units(blocks: list) -> Iterator[tuple[str, int, int]]:
    """Flatten blocks into packable units, each <= CHUNK_MAX_WORDS *or* an
    indivisible single sentence. An oversized paragraph is broken at sentence
    boundaries only; a lone sentence over the ceiling is emitted whole (we never
    cut a sentence)."""
    for text, start, end in blocks:
        if _approx_tokens(text) <= CHUNK_MAX_WORDS:
            yield text, start, end
        else:
            for sent in _SENT_SPLIT.split(text):
                sent = sent.strip()
                if sent:
                    yield sent, start, end   # line span approximated to the paragraph


def _chunk_blocks(blocks: list) -> Iterator[tuple[str, int, int]]:
    """Pack a section's blocks into ~CHUNK_TARGET_WORDS chunks with sentence
    overlap. Yields (raw_text, start_line, end_line). Only reached for sections
    that exceed CHUNK_MAX_WORDS — smaller sections are emitted whole upstream."""
    buf: list[str] = []
    btok = 0
    bstart = bend = None
    prev = None  # last emitted chunk text, for overlap seeding

    def close():
        nonlocal buf, btok, bstart, bend, prev
        if buf:
            prev = "\n".join(buf)
            yield_val = (prev, bstart, bend)
            buf, btok, bstart, bend = [], 0, None, None
            return yield_val
        return None

    for text, start, end in _units(blocks):
        t = _approx_tokens(text)
        if buf and btok + t > CHUNK_TARGET_WORDS:
            out = close()
            if out:
                yield out
        if not buf:
            tail = _overlap_tail(prev) if prev else ""
            if tail:
                buf.append(tail); btok += _approx_tokens(tail)
            bstart, bend = start, end   # span tracks the real units, not the overlap
        buf.append(text); btok += t; bend = end
    out = close()
    if out:
        yield out


def _emit_sections(path: pathlib.Path, sections: list, markdown: bool) -> Iterator[dict]:
    """Turn sections into chunk records. A section at/under CHUNK_MAX_WORDS is kept
    WHOLE (one chunk) so a coherent section is never split just to hit the target;
    larger sections are recursively split at sentence boundaries with overlap."""
    cidx = 0
    for sec_title, sub_title, blocks in sections:
        if not blocks:
            continue
        total = sum(_approx_tokens(b[0]) for b in blocks)
        if total <= CHUNK_MAX_WORDS:
            pieces = [("\n".join(b[0] for b in blocks), blocks[0][1], blocks[-1][2])]
        else:
            pieces = _chunk_blocks(blocks)
        for raw, start, end in pieces:
            raw = raw.strip()
            if not raw:
                continue
            clean = _strip_md(raw) if markdown else raw
            yield {
                "id": point_id(path.name, sec_title or "_", cidx),
                "text": raw,          # verbatim -> returned to the LLM
                "text_clean": clean,  # decoration-stripped -> embedding input only
                "payload": {
                    "source_file": path.name,
                    "section_title": sec_title,
                    "subsection_title": sub_title,
                    "heading": sec_title,            # kept for current retrieval's `location`
                    "chunk_index": cidx,
                    "chunk_start_line": start,
                    "chunk_end_line": end,
                    "kind": "doc_text",
                    "lang": detect_lang(clean),
                    "text": raw,
                    "text_clean": clean,
                },
            }
            cidx += 1


def _sections_from_lines(lines: list, markdown: bool) -> list:
    """Parse lines into sections with heading titles and line-spanned blocks."""
    sections: list = []
    sec = sub = ""
    blocks: list = []
    buf: list[str] = []
    buf_start = None

    def flush_para(end_line):
        nonlocal buf, buf_start
        if buf:
            blocks.append((" ".join(buf), buf_start, end_line))
            buf, buf_start = [], None

    def flush_section():
        nonlocal blocks
        if blocks:
            sections.append((sec, sub, blocks)); blocks = []

    for i, ln in enumerate(lines, start=1):
        s = ln.rstrip("\n")
        m = _MD_HEADING.match(s) if markdown else None
        if m:
            flush_para(i - 1)
            flush_section()
            sec, sub = _heading_titles(len(m.group(1)), m.group(2).strip(), sec, sub)
        elif s.strip() == "":
            flush_para(i - 1)
        else:
            if buf_start is None:
                buf_start = i
            buf.append(s.strip())
    flush_para(len(lines))
    flush_section()
    return sections


def iter_text(path: pathlib.Path, markdown: bool) -> Iterator[dict]:
    raw = path.read_text(encoding="utf-8", errors="replace")
    sections = _sections_from_lines(raw.splitlines(), markdown)
    yield from _emit_sections(path, sections, markdown)


def iter_pdf(path: pathlib.Path) -> Iterator[dict]:
    if pdfplumber is None:
        log.error("pdfplumber not installed; cannot read %s", path)
        return
    page_texts = []
    with pdfplumber.open(str(path)) as pdf:
        for pg in pdf.pages:
            t = pg.extract_text() or ""
            if t.strip():
                page_texts.append(t)
    if not page_texts:
        log.warning("%s: no extractable text layer (scanned?) — skipping", path.name)
        return
    lines = "\n\n".join(page_texts).splitlines()
    sections = _sections_from_lines(lines, markdown=False)
    yield from _emit_sections(path, sections, markdown=False)


def _docx_heading_level(p) -> int:
    """Heading level for a docx paragraph, or 0 if it's body text.

    Two signals, because this corpus's headings are often *bold body text* rather
    than real Heading styles (the reason whole sections used to collapse into one
    giant chunk):
      1. a real 'Heading N' / 'Title' style  -> level N (Title -> 1);
      2. a short, standalone, fully-bold line -> level 1 (bold-run fallback).
    """
    name = (p.style.name or "").lower() if p.style is not None else ""
    if name.startswith("heading") or name == "title":
        m = re.search(r"(\d+)", name)
        return int(m.group(1)) if m else 1
    txt = p.text.strip()
    if txt and _approx_tokens(txt) <= 12 and len(txt) <= 80:
        runs = [r for r in p.runs if r.text.strip()]
        style_bold = bool(p.style is not None and p.style.font is not None and p.style.font.bold)
        if runs and all(bool(r.bold) or style_bold for r in runs):
            return 1
    return 0


def iter_docx(path: pathlib.Path) -> Iterator[dict]:
    if Document is None:
        log.error("python-docx not installed; cannot read %s", path)
        return
    doc = Document(str(path))

    for ti, table in enumerate(doc.tables):
        rows = table.rows
        if not rows:
            continue
        keys = {i: c.text.strip() for i, c in enumerate(rows[0].cells) if c.text.strip()}
        for ridx, row in enumerate(rows[1:], start=1):
            values = {i: row.cells[i].text for i in keys if i < len(row.cells)}
            text = _serialize_kv(keys, values, f"table{ti}")
            if not text:
                continue
            yield {
                "id": point_id(path.name, f"table{ti}", ridx),
                "text": text,
                "payload": {
                    "source_file": path.name, "sheet": f"table{ti}", "row_index": ridx,
                    "kind": "table_row", "lang": "n/a", "text": text,
                },
            }

    sections: list = []
    sec = sub = ""
    blocks: list = []

    def flush_section():
        nonlocal blocks
        if blocks:
            sections.append((sec, sub, blocks)); blocks = []

    # Paragraph ordinal (1-based over ALL paragraphs incl. blanks) is the
    # docx "line" proxy — Word has no source lines, so chunk_start/end_line cite
    # paragraph position instead.
    for i, p in enumerate(doc.paragraphs, start=1):
        txt = p.text.strip()
        if not txt:
            continue
        lvl = _docx_heading_level(p)
        if lvl:
            flush_section()
            sec, sub = _heading_titles(lvl, txt, sec, sub)
        else:
            blocks.append((txt, i, i))
    flush_section()
    yield from _emit_sections(path, sections, markdown=False)


# --- routing ----------------------------------------------------------------
def iter_file(path: pathlib.Path) -> Iterator[dict]:
    ext = path.suffix.lower()
    if ext in (".xlsx", ".xls"):
        yield from iter_xlsx(path)
    elif ext == ".csv":
        yield from iter_csv(path)
    elif ext == ".docx":
        yield from iter_docx(path)
    elif ext == ".txt":
        yield from iter_text(path, markdown=False)
    elif ext in (".md", ".markdown"):
        yield from iter_text(path, markdown=True)
    elif ext == ".pdf":
        yield from iter_pdf(path)
    else:
        log.warning("skipping unsupported file: %s", path.name)


# ============================================================================
# COLLECTION + ALIAS (ADR-0006)
# ============================================================================
def qclient() -> QdrantClient:
    return QdrantClient(url=QDRANT_URL)


def _create_physical(qc: QdrantClient, name: str):
    """Create a physical collection with the dense+sparse schema (idempotent)."""
    if not qc.collection_exists(name):
        qc.create_collection(
            name,
            vectors_config={"dense": qm.VectorParams(size=DENSE_DIM, distance=qm.Distance.COSINE)},
            sparse_vectors_config={"sparse": qm.SparseVectorParams()},
        )
        log.info("created physical collection %s (dense %d + sparse)", name, DENSE_DIM)


def resolve_alias(qc: QdrantClient, alias: str) -> str | None:
    """Physical collection `alias` points to, or None if `alias` is not an alias."""
    try:
        for a in qc.get_aliases().aliases:
            if a.alias_name == alias:
                return a.collection_name
    except Exception:
        pass
    return None


def live_collection(qc: QdrantClient, alias: str) -> str:
    """
    The collection incremental ops target. Prefer the alias' physical collection;
    tolerate the legacy state where `alias` is still a plain physical collection
    (pre-ADR-0006, before the first --recreate migrates it behind an alias).
    """
    phys = resolve_alias(qc, alias)
    if phys:
        return phys
    if qc.collection_exists(alias):
        return alias  # legacy physical collection named like the alias
    raise SystemExit(f"no collection or alias '{alias}' exists — run `--recreate` first")


# ============================================================================
# EMBED + UPSERT (uuid4 point IDs; frozen embed() reused from above)
# ============================================================================
def _context_prefix(payload: dict) -> str:
    """A retrieval-only lead-in prepended to a chunk's EMBEDDING input (never to the
    verbatim `text` the model quotes). It surfaces the two things a non-expert author
    can reliably get right — the FILE NAME and the SHEET/SECTION — as plain tokens, so
    a topic query ('firewall rules') matches even when the row content itself (IPs,
    ports, TCP) carries no such vocabulary. Separators are normalised to spaces so
    'EPC25-FW-rules-...' tokenizes into FW / rules / ..."""
    stem = re.sub(r"[-_]+", " ", pathlib.Path(payload.get("source_file", "")).stem).strip()
    loc = payload.get("sheet") or payload.get("section_title") or ""
    parts = [f"file: {stem}"] if stem else []
    if loc:
        parts.append(str(loc))
    return (" | ".join(parts) + " | ") if parts else ""


def _upsert_records(qc: QdrantClient, collection: str, records: list[dict]) -> list[str]:
    """Embed a batch of chunk records and upsert with fresh uuid4 IDs.

    The record's own `id` (a deterministic hash from the frozen parser) is IGNORED:
    identity now lives in the checksum manifest, not the point ID (ADR-0006). This
    is what makes blind delete-by-recorded-UUID safe — old and new versions of a
    file never share an ID.
    """
    if not records:
        return []
    # Embed the decoration-stripped text_clean (prose); table rows have no clean
    # variant, so fall back to their verbatim text. Each embed input is prefixed with
    # a filename + sheet/section context lead-in (_context_prefix) so topic/filename
    # queries retrieve. text_clean stores EXACTLY what was embedded; the verbatim
    # `text` returned to the model is untouched (no lead-in).
    emb_texts = [_context_prefix(r["payload"]) + (r.get("text_clean") or r["text"]) for r in records]
    dense, sparse = embed(emb_texts, kind="passage")
    points, ids = [], []
    for r, et, dv, sv in zip(records, emb_texts, dense, sparse):
        r["payload"]["text_clean"] = et
        pid = str(uuid.uuid4())
        ids.append(pid)
        points.append(qm.PointStruct(
            id=pid,
            vector={"dense": dv, "sparse": to_sparse_vector(sv)},
            payload=r["payload"],
        ))
    qc.upsert(collection, points)
    return ids


def ingest_file(qc: QdrantClient, collection: str, path: pathlib.Path) -> list[str]:
    """Parse+embed+upsert one file. Returns all point UUIDs it produced."""
    ids: list[str] = []
    batch: list[dict] = []
    for rec in iter_file(path):
        batch.append(rec)
        if len(batch) >= BATCH:
            ids += _upsert_records(qc, collection, batch); batch = []
    if batch:
        ids += _upsert_records(qc, collection, batch)
    return ids


def _refresh_source_file(qc: QdrantClient, collection: str, uuids: list[str], name: str):
    """Same bytes under a new filename -> keep payload source_file honest, cheaply."""
    if uuids:
        qc.set_payload(collection, payload={"source_file": name},
                       points=qm.PointIdsList(points=uuids))


# ============================================================================
# DISK SCAN + CHECKSUM
# ============================================================================
def file_checksum(path: pathlib.Path) -> str:
    h = hashlib.blake2b(digest_size=32)
    with path.open("rb") as fh:
        for block in iter(lambda: fh.read(1 << 20), b""):
            h.update(block)
    return h.hexdigest()


def scan_disk(root: pathlib.Path) -> dict[str, str]:
    """Map content checksum -> relative path for every file under root.

    Identical-content files collapse to one checksum (last path wins) — intended
    dedup (see corpus_manifest)."""
    out: dict[str, str] = {}
    for f in sorted(root.rglob("*")):
        if f.is_file():
            out[file_checksum(f)] = str(f.relative_to(root))
    return out


# ============================================================================
# SYNC PASSES
# ============================================================================
def incremental(qc: QdrantClient, root: pathlib.Path, man: Manifest,
                grace: float | None) -> dict:
    """
    In-place sync of the live collection against disk.

    grace is None -> MANUAL: departed checksums pruned immediately.
    grace is a num -> WATCH:  departed checksums only *marked*; the reaper deletes
                     them once absent >= grace seconds (rides out atomic saves /
                     rsync windows). Reappearance before grace cancels the delete.

    Ordering is embed-before-prune, so no file's data is ever missing mid-pass.
    """
    collection = live_collection(qc, ALIAS)
    disk = scan_disk(root)                     # checksum -> relpath
    disk_ck = set(disk)
    db = man.all()                             # checksum -> {uuids, path, departed_at}
    db_ck = set(db)

    added = renamed = pruned = 0

    # 1. ADD new content (new UUIDs) — before any prune
    for ck in disk_ck - db_ck:
        ids = ingest_file(qc, collection, root / disk[ck])
        man.put(ck, ids, disk[ck])
        added += 1

    # 2. RECONCILE still-present content: cancel stale departals, refresh renames
    for ck in disk_ck & db_ck:
        rec = db[ck]
        if rec["departed_at"] is not None:
            man.clear_departed(ck)
        if rec["path"] != disk[ck]:
            _refresh_source_file(qc, collection, rec["uuids"], pathlib.Path(disk[ck]).name)
            man.set_path(ck, disk[ck])
            renamed += 1

    # 3. DEPARTED content
    now = time.time()
    for ck in db_ck - disk_ck:
        if grace is None:
            qc.delete(collection, points_selector=qm.PointIdsList(points=man.delete(ck)))
            pruned += 1
        else:
            man.mark_departed(ck, now)

    # 3b. WATCH reaper: execute grace-expired deletes
    if grace is not None:
        for ck, uuids in man.departed_due(now, grace):
            qc.delete(collection, points_selector=qm.PointIdsList(points=uuids))
            man.delete(ck)
            pruned += 1

    log.info("incremental: +%d added, %d renamed, -%d pruned (%d files) -> %s",
             added, renamed, pruned, len(disk_ck), collection)
    return {"collection": collection, "added": added, "renamed": renamed,
            "pruned": pruned, "on_disk": len(disk_ck)}


def recreate(qc: QdrantClient, root: pathlib.Path, man: Manifest) -> dict:
    """
    Full rebuild with ZERO retrieval blackout via alias switch: build corpus_<ts>
    alongside the live one, embed everything, then atomically repoint ALIAS and
    drop the old physical collection. The ONLY path using dual-collection.
    """
    # timestamp is for humans; the uuid suffix guarantees a UNIQUE physical name
    # even on sub-second repeated --recreate runs (strftime is 1s-resolution).
    # Reusing an existing collection would orphan the prior run's points and break
    # the build-alongside/zero-blackout guarantee.
    ts = time.strftime("%Y%m%d-%H%M%S")
    newcol = f"{ALIAS}_{ts}_{uuid.uuid4().hex[:8]}"
    _create_physical(qc, newcol)

    disk = scan_disk(root)
    rows: list[tuple[str, list[str], str]] = []
    for ck, rel in disk.items():
        ids = ingest_file(qc, newcol, root / rel)
        rows.append((ck, ids, rel))
    man.replace_all(rows)

    old = resolve_alias(qc, ALIAS)
    if old is None and qc.collection_exists(ALIAS):
        # ONE-TIME legacy migration: ALIAS is still a physical collection. Can't
        # alias over an existing collection name, so drop it then alias. This is
        # the sole, momentary blackout in the whole design — it happens once, ever.
        qc.delete_collection(ALIAS)
        qc.update_collection_aliases(change_aliases_operations=[
            qm.CreateAliasOperation(create_alias=qm.CreateAlias(
                collection_name=newcol, alias_name=ALIAS)),
        ])
        log.info("recreate: migrated legacy physical '%s' -> alias -> %s", ALIAS, newcol)
    else:
        ops = []
        if old:
            ops.append(qm.DeleteAliasOperation(delete_alias=qm.DeleteAlias(alias_name=ALIAS)))
        ops.append(qm.CreateAliasOperation(create_alias=qm.CreateAlias(
            collection_name=newcol, alias_name=ALIAS)))
        qc.update_collection_aliases(change_aliases_operations=ops)  # atomic
        if old and old != newcol:
            qc.delete_collection(old)
        log.info("recreate: alias '%s' -> %s (dropped %s)", ALIAS, newcol, old or "-")

    points = sum(len(u) for _, u, _ in rows)
    log.info("recreate: %d files -> %d points in %s", len(disk), points, newcol)
    return {"collection": newcol, "files": len(disk), "points": points}


# ============================================================================
# WATCH MODE
# ============================================================================
def watch(root: pathlib.Path, man: Manifest):
    """
    Autonomous, debounced watcher (ADR-0006 §watch). Differs from a naive
    "inotify -> ingest": debounce/quiescence so we never read a half-written file
    and coalesce event storms; add/update prompt but deletes GRACE-gated; a full
    scan every SCAN_INTERVAL as the missed-event backstop; each pass under a
    non-blocking lock so it yields to a manual run.
    """
    from watchdog.observers import Observer
    from watchdog.observers.polling import PollingObserver
    from watchdog.events import FileSystemEventHandler

    state = {"dirty": True, "last_event": 0.0, "last_scan": 0.0}  # dirty -> initial pass

    # React only to content-mutating events. Every pass OPENS+READS each corpus
    # file to checksum it, which the observer sees as IN_OPEN / IN_CLOSE_NOWRITE
    # ('opened' / 'closed_no_write'). Reacting to those read-only events would let
    # the watcher's own hashing re-arm the debounce and spin a pass every DEBOUNCE
    # seconds forever (a self-feeding loop). Keep writes/creates/deletes/renames
    # and close-after-write; ignore bare reads.
    _READONLY_EVENTS = ("opened", "closed_no_write")

    class _H(FileSystemEventHandler):
        def on_any_event(self, event):
            if getattr(event, "is_directory", False):
                return
            if event.event_type in _READONLY_EVENTS:
                return
            state["dirty"] = True
            state["last_event"] = time.time()

    obs = PollingObserver() if WATCH_POLL else Observer()
    obs.schedule(_H(), str(root), recursive=True)
    obs.start()
    log.info("watching %s (debounce %ss, delete-grace %ss, full-scan %ss, %s)",
             root, DEBOUNCE, GRACE, SCAN_INTERVAL, "polling" if WATCH_POLL else "inotify")

    def _pass():
        try:
            with pass_lock(man.db_path, blocking=False):
                incremental(qclient(), root, man, grace=GRACE)
        except Locked:
            log.info("pass skipped: a manual ingest holds the lock; retry next tick")
        except Exception:
            log.exception("watch pass failed; continuing")

    try:
        while True:
            time.sleep(1.0)
            now = time.time()
            due_event = state["dirty"] and (now - state["last_event"]) >= DEBOUNCE
            due_scan = (now - state["last_scan"]) >= SCAN_INTERVAL
            if due_event or due_scan:
                state["dirty"] = False
                state["last_scan"] = now
                _pass()
    except KeyboardInterrupt:
        pass
    finally:
        obs.stop()
        obs.join()
        log.info("watcher stopped")


# ============================================================================
# MAIN
# ============================================================================
def main():
    ap = argparse.ArgumentParser(description="ragfarm corpus ingester (ADR-0006)")
    ap.add_argument("--corpus", default=str(ROOT), help="corpus root (default: $CORPUS_PATH)")
    mode = ap.add_mutually_exclusive_group()
    mode.add_argument("--recreate", action="store_true",
                      help="full rebuild + atomic alias switch (manifest-blind)")
    mode.add_argument("--watch", action="store_true",
                      help="run as a debounced autonomous watcher")
    args = ap.parse_args()

    root = pathlib.Path(args.corpus)
    if not root.exists():
        log.error("corpus root does not exist: %s", root); sys.exit(1)

    man = Manifest(MANIFEST_DB)
    try:
        if args.watch:
            watch(root, man)                                    # own per-pass locking
        elif args.recreate:
            with pass_lock(man.db_path, blocking=True):
                recreate(qclient(), root, man)
        else:
            with pass_lock(man.db_path, blocking=True):
                incremental(qclient(), root, man, grace=None)   # manual: prune now
    finally:
        man.close()


if __name__ == "__main__":
    main()
